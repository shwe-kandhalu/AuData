"""Build a downloadable Word (.docx) audit report that synthesizes every
detector's findings for a paper: Statistical Recompute (p-values), Meta-analysis,
Numerical Consistency, Image Forensics, Methods<->Claims, and Reference Integrity.

Reads the persisted per-paper audit stages (the same data the UI Audit Report
shows) and renders a clean, professional Word document.
"""

from __future__ import annotations

import io
from datetime import datetime
from typing import Any, Dict, List

from . import storage

_SEV_RANK = {"high": 4, "medium": 3, "moderate": 3, "low": 2, "info": 1, "none": 0}


def _sev(x: Any) -> str:
    s = (str(x) if x else "medium").lower()
    return "medium" if s == "moderate" else s


def _norm(key: str, d: Any) -> Dict[str, Any]:
    """Normalize one stage's data -> {ran, total, flagged, items:[{title,severity,detail}], note}."""
    if not d:
        return {"ran": False}
    items: List[Dict[str, str]] = []
    if key == "references":
        res = d.get("results", []) or []
        fl = [r for r in res if r.get("status") == "flagged"]
        for r in fl:
            title = (r.get("matched") or {}).get("title") or (r.get("input") or {}).get("raw") or f"Reference {r.get('number','')}"
            items.append({"title": title, "severity": _sev(r.get("severity")),
                          "detail": "; ".join(i.get("label", "") for i in (r.get("issues") or []))})
        return {"ran": True, "total": (d.get("summary") or {}).get("total", len(res)),
                "flagged": (d.get("summary") or {}).get("flagged", len(fl)), "items": items}
    if key == "methods":
        res = d.get("results", []) or []
        fl = [r for r in res if r.get("status") == "flagged"]
        for r in fl:
            items.append({"title": r.get("claim", ""), "severity": _sev(r.get("severity")),
                          "detail": " - ".join(x for x in (r.get("issue_type"), r.get("reasoning")) if x)})
        return {"ran": True, "total": (d.get("summary") or {}).get("total", len(res)),
                "flagged": (d.get("summary") or {}).get("flagged", len(fl)), "items": items}
    if key == "statcheck":
        f = d.get("findings", []) or []
        fl = [x for x in f if x.get("status") == "mismatch"]
        for x in fl:
            items.append({"title": x.get("claim", ""), "severity": "high",
                          "detail": x.get("note", "Reported p-value does not match the recomputed value.")})
        return {"ran": True, "total": d.get("claim_count", len(f)),
                "flagged": d.get("mismatch_count", len(fl)), "items": items}
    if key == "numerical":
        flags = d.get("flags", []) or []
        for x in flags:
            items.append({"title": x.get("description") or x.get("type") or "Inconsistency",
                          "severity": _sev(x.get("severity")),
                          "detail": f'"{x.get("excerpt")}"' if x.get("excerpt") else (x.get("type") or "")})
        return {"ran": True, "total": len(flags), "flagged": len(flags), "items": items}
    if key == "images":
        summ = d.get("summary") or {}
        report = d.get("report") or {}
        for fnd in (report.get("cross_paper_findings") or []):
            score = fnd.get("similarity_score")
            items.append({"title": f"{fnd.get('flag_type','Figure reuse')}: {fnd.get('target_figure','')} ~ {fnd.get('candidate_figure','')}",
                          "severity": _sev(fnd.get("severity")),
                          "detail": f"cross-paper similarity {score:.2f}" if isinstance(score, (int, float)) else f"cross-paper similarity {score}"})
        for r in (report.get("figure_forensics") or []):
            fig = f"page {r.get('metadata',{}).get('page')}" if r.get("metadata", {}).get("page") else "a figure"
            cm = r.get("copy_move_result") or {}
            sp = r.get("splice_result") or {}
            if cm.get("severity") and cm["severity"] not in ("low", "none"):
                items.append({"title": f"Copy-move in {fig}", "severity": _sev(cm["severity"]), "detail": "Cloned region detected within the figure."})
            if sp.get("severity") and sp["severity"] not in ("low", "none"):
                items.append({"title": f"Splice boundary in {fig}", "severity": _sev(sp["severity"]), "detail": "Possible splice / edited boundary."})
            ai = r.get("ai_generated_score")
            if isinstance(ai, (int, float)) and ai >= 0.7:
                items.append({"title": f"Possibly AI-generated ({fig})", "severity": "medium", "detail": f"heuristic score {ai:.2f}"})
            vlm = r.get("vlm_result") or {}
            if vlm.get("verdict") and vlm["verdict"] != "clean" and (vlm.get("confidence") or 0) >= 0.5:
                label = "possibly AI-generated" if vlm["verdict"] == "ai_generated" else "manipulation suspected"
                items.append({"title": f"Vision model: {label} ({fig})", "severity": "medium", "detail": vlm.get("reason", "")})
        return {"ran": True, "total": summ.get("total_images", report.get("num_target_figures", 0)),
                "flagged": summ.get("flagged", len(items)), "items": items}
    if key == "meta":
        if not d.get("detected"):
            return {"ran": True, "total": 0, "flagged": 0, "items": [], "note": "No meta-analysis detected in this paper."}
        disc = d.get("verdict") == "discrepancy"
        if disc:
            items.append({"title": f"Pooled {d.get('measure','')} discrepancy", "severity": _sev(d.get("severity")),
                          "detail": d.get("explanation", "")})
        return {"ran": True, "total": 1, "flagged": 1 if disc else 0, "items": items,
                "note": None if disc else d.get("explanation")}
    return {"ran": False}


_SECTIONS = [
    ("statcheck", "Statistical Recompute (p-values)"),
    ("meta", "Meta-analysis recreation"),
    ("numerical", "Numerical Consistency"),
    ("images", "Image Forensics"),
    ("methods", "Methods / Claims"),
    ("references", "Reference Integrity"),
]

# Restrained palette: structure in grayscale; severity only as small colored text.
_NAVY = (0x1F, 0x3A, 0x5F)
_GRAY = (0x60, 0x60, 0x60)
_HEADER_FILL = "F2F2F2"
_SEV_TXT = {
    "high":   (0x9C, 0x1F, 0x1F),
    "medium": (0x8A, 0x52, 0x00),
    "low":    (0x1F, 0x5C, 0x8B),
    "info":   (0x60, 0x60, 0x60),
    "none":   (0x1E, 0x7A, 0x3A),
}


def _shade(cell, fill_hex: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _hrule(paragraph) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "BFBFBF")
    pbdr.append(bottom)
    pPr.append(pbdr)


def build_docx(paper_id: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches

    paper = storage.get_paper(paper_id) or {}
    audits = storage.get_paper_audits(paper_id) or {}
    sections = [(label, _norm(key, audits.get(key))) for key, label in _SECTIONS]
    run_sections = [(l, n) for l, n in sections if n.get("ran")]
    total_flagged = sum(n.get("flagged", 0) for _, n in run_sections)
    by_sev: Dict[str, int] = {}
    for _, n in run_sections:
        for it in n.get("items", []):
            by_sev[it["severity"]] = by_sev.get(it["severity"], 0) + 1
    clean = bool(run_sections) and total_flagged == 0

    doc = Document()
    try:
        doc.styles["Normal"].font.size = Pt(10.5)
    except Exception:
        pass

    title = doc.add_heading("Research Integrity Audit Report", level=0)
    for r in title.runs:
        r.font.color.rgb = RGBColor(*_NAVY)

    tp = doc.add_paragraph()
    tr = tp.add_run(paper.get("title") or paper_id)
    tr.bold = True
    tr.font.size = Pt(13)
    meta_line = " · ".join(str(x) for x in (paper.get("authors"), paper.get("year"), paper.get("container")) if x)
    if meta_line:
        mp = doc.add_paragraph()
        mr = mp.add_run(meta_line)
        mr.font.size = Pt(9)
        mr.font.color.rgb = RGBColor(*_GRAY)
    gp = doc.add_paragraph()
    gr = gp.add_run(f"Generated {datetime.now().strftime('%B %d, %Y at %H:%M')}")
    gr.italic = True
    gr.font.size = Pt(8)
    gr.font.color.rgb = RGBColor(*_GRAY)
    _hrule(gp)

    # ── Summary ──
    doc.add_heading("Summary", level=2)
    sp = doc.add_paragraph()
    if run_sections:
        verdict = ("No integrity issues were detected" if clean
                   else f"{total_flagged} potential issue{'' if total_flagged == 1 else 's'} flagged for review")
        vr = sp.add_run(verdict)
        vr.bold = True
        vr.font.color.rgb = RGBColor(*(_SEV_TXT["none"] if clean else _SEV_TXT["high"]))
        sp.add_run(f" across {len(run_sections)} of {len(_SECTIONS)} detectors run.")
        if by_sev:
            bp = doc.add_paragraph()
            bl = bp.add_run("Severity breakdown:  ")
            bl.bold = True
            bl.font.size = Pt(9)
            parts = [f"{by_sev[k]} {k}" for k in ("high", "medium", "low", "info") if by_sev.get(k)]
            br = bp.add_run(", ".join(parts) or "none")
            br.font.size = Pt(9)
            br.font.color.rgb = RGBColor(*_GRAY)
    else:
        sp.add_run("No detectors have been run for this paper yet.")

    # detectors overview table
    dt = doc.add_table(rows=1, cols=3)
    dt.style = "Table Grid"
    for i, h in enumerate(("Detector", "Result", "Findings")):
        c = dt.rows[0].cells[i]
        _shade(c, _HEADER_FILL)
        rr = c.paragraphs[0].add_run(h)
        rr.bold = True
        rr.font.size = Pt(9)
    for label, n in sections:
        cells = dt.add_row().cells
        cells[0].paragraphs[0].add_run(label).font.size = Pt(10)
        ran = n.get("ran")
        flagged = n.get("flagged", 0)
        if not ran:
            txt, tone = "Not run", "info"
        elif flagged:
            txt, tone = "Flagged", "high"
        else:
            txt, tone = "Clean", "none"
        res = cells[1].paragraphs[0].add_run(("✓ " if tone == "none" else "") + txt)
        res.bold = True
        res.font.size = Pt(9.5)
        res.font.color.rgb = RGBColor(*_SEV_TXT[tone])
        cnt_txt = (f"{flagged} flagged of {n.get('total')}" if ran and n.get("total") else (f"{flagged} flagged" if ran else "—"))
        cr = cells[2].paragraphs[0].add_run(cnt_txt)
        cr.font.size = Pt(9)
        cr.font.color.rgb = RGBColor(*_GRAY)
    for row in dt.rows:
        row.cells[0].width = Inches(3.1)
        row.cells[1].width = Inches(1.3)
        row.cells[2].width = Inches(2.1)

    # ── Detailed findings ──
    if any(n.get("flagged") for _, n in run_sections):
        doc.add_heading("Detailed findings", level=2)
        for label, n in sections:
            if not n.get("ran") or not n.get("flagged"):
                continue
            doc.add_heading(label, level=3)
            items = sorted(n.get("items", []), key=lambda it: -_SEV_RANK.get(it["severity"], 0))
            ft = doc.add_table(rows=1, cols=2)
            ft.style = "Table Grid"
            for i, h in enumerate(("Severity", "Finding")):
                c = ft.rows[0].cells[i]
                _shade(c, _HEADER_FILL)
                rr = c.paragraphs[0].add_run(h)
                rr.bold = True
                rr.font.size = Pt(9)
            for it in items:
                cells = ft.add_row().cells
                sv = cells[0].paragraphs[0].add_run(it["severity"].upper())
                sv.bold = True
                sv.font.size = Pt(8.5)
                sv.font.color.rgb = RGBColor(*_SEV_TXT.get(it["severity"], _SEV_TXT["info"]))
                tr2 = cells[1].paragraphs[0].add_run(it["title"])
                tr2.bold = True
                tr2.font.size = Pt(9.5)
                if it.get("detail"):
                    dp = cells[1].add_paragraph()
                    dr = dp.add_run(it["detail"])
                    dr.font.size = Pt(9)
                    dr.font.color.rgb = RGBColor(*_GRAY)
            for row in ft.rows:
                row.cells[0].width = Inches(0.9)
                row.cells[1].width = Inches(5.6)

    # ── Passed checks ──
    passed = [label for label, n in sections if n.get("ran") and not n.get("flagged")]
    if passed:
        doc.add_heading("Passed checks", level=2)
        pp = doc.add_paragraph()
        for i, label in enumerate(passed):
            mark = pp.add_run(("     " if i else "") + "✓ ")
            mark.bold = True
            mark.font.color.rgb = RGBColor(*_SEV_TXT["none"])
            t = pp.add_run(label)
            t.font.size = Pt(10)

    doc.add_paragraph()
    foot = doc.add_paragraph()
    fr = foot.add_run("Generated by AuData. Flags are decision support for a human reviewer, not determinations of misconduct.")
    fr.italic = True
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(*_GRAY)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
